"""
CLI & Web Generator: Builds downloadable PDF, DOCX, and ZIP archives for all 10 CVs and 10 JDs.
"""

import os
import io
import zipfile
import docx
from core.pdf_builder import generate_pdf_report
from core.sample_data import SAMPLE_JDS, SAMPLE_RESUMES


def create_docx(title: str, text: str) -> bytes:
    """Generates a styled DOCX document in-memory."""
    doc = docx.Document()
    doc.add_heading(title, level=1)
    for paragraph in text.split("\n\n"):
        clean_p = paragraph.strip()
        if clean_p:
            doc.add_paragraph(clean_p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_all_cvs_and_jds_zip() -> bytes:
    """Creates a consolidated ZIP containing all 10 CVs and 10 JDs in both PDF and DOCX formats."""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        # 1. Package 10 CVs (PDF & DOCX)
        for idx, cv in enumerate(SAMPLE_RESUMES, start=1):
            base_name = f"CV_{idx:02d}_{cv['candidate_name'].replace(' ', '_')}"
            # PDF
            pdf_data = generate_pdf_report(cv["raw_content"], title=f"Curriculum Vitae - {cv['candidate_name']}")
            zip_file.writestr(f"Resumes_CVs/PDF/{base_name}.pdf", pdf_data)
            # DOCX
            docx_data = create_docx(f"Curriculum Vitae - {cv['candidate_name']}", cv["raw_content"])
            zip_file.writestr(f"Resumes_CVs/DOCX/{base_name}.docx", docx_data)

        # 2. Package 10 JDs (PDF & DOCX)
        for idx, jd in enumerate(SAMPLE_JDS, start=1):
            base_name = f"JD_{idx:02d}_{jd['title'].replace(' ', '_').replace('/', '_')}"
            # PDF
            pdf_data = generate_pdf_report(jd["content"], title=f"Job Specification - {jd['title']}")
            zip_file.writestr(f"Job_Descriptions/PDF/{base_name}.pdf", pdf_data)
            # DOCX
            docx_data = create_docx(f"Job Specification - {jd['title']}", jd["content"])
            zip_file.writestr(f"Job_Descriptions/DOCX/{base_name}.docx", docx_data)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def export_files_to_disk(output_dir: str = "./sample_downloads"):
    """Saves all 10 CVs and 10 JDs as physical files in the specified directory."""
    os.makedirs(f"{output_dir}/Resumes_PDF", exist_ok=True)
    os.makedirs(f"{output_dir}/Resumes_DOCX", exist_ok=True)
    os.makedirs(f"{output_dir}/JDs_PDF", exist_ok=True)
    os.makedirs(f"{output_dir}/JDs_DOCX", exist_ok=True)

    print(f"Generating documents into '{output_dir}'...")

    for idx, cv in enumerate(SAMPLE_RESUMES, start=1):
        name = f"CV_{idx:02d}_{cv['candidate_name'].replace(' ', '_')}"
        pdf_b = generate_pdf_report(cv["raw_content"], title=f"CV - {cv['candidate_name']}")
        docx_b = create_docx(f"CV - {cv['candidate_name']}", cv["raw_content"])
        
        with open(f"{output_dir}/Resumes_PDF/{name}.pdf", "wb") as f:
            f.write(pdf_b)
        with open(f"{output_dir}/Resumes_DOCX/{name}.docx", "wb") as f:
            f.write(docx_b)

    for idx, jd in enumerate(SAMPLE_JDS, start=1):
        name = f"JD_{idx:02d}_{jd['title'].replace(' ', '_').replace('/', '_')}"
        pdf_b = generate_pdf_report(jd["content"], title=f"JD - {jd['title']}")
        docx_b = create_docx(f"JD - {jd['title']}", jd["content"])

        with open(f"{output_dir}/JDs_PDF/{name}.pdf", "wb") as f:
            f.write(pdf_b)
        with open(f"{output_dir}/JDs_DOCX/{name}.docx", "wb") as f:
            f.write(docx_b)

    print(f"[✓] Successfully generated 10 CVs and 10 JDs in PDF and DOCX formats under '{output_dir}'.")


if __name__ == "__main__":
    export_files_to_disk()
